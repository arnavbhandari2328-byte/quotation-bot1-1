import os
import base64
from datetime import datetime

from flask import Flask, request, jsonify
from google.oauth2.credentials import Credentials
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError

from email.mime.base import MIMEBase
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email import encoders

from docx import Document

app = Flask(__name__)

# ------------ ENV needed ------------
REQUIRED_ENV = [
    "GMAIL_CLIENT_ID",
    "GMAIL_CLIENT_SECRET",
    "GMAIL_REFRESH_TOKEN",
    "GMAIL_SENDER",
]

def have_gmail_env():
    return all(os.getenv(k) for k in REQUIRED_ENV)

def gmail_service():
    creds = Credentials(
        token=None,
        refresh_token=os.environ["GMAIL_REFRESH_TOKEN"],
        token_uri="https://oauth2.googleapis.com/token",
        client_id=os.environ["GMAIL_CLIENT_ID"],
        client_secret=os.environ["GMAIL_CLIENT_SECRET"],
        scopes=["https://www.googleapis.com/auth/gmail.send"],
    )
    return build("gmail", "v1", credentials=creds, cache_discovery=False)

# ------------ Template helpers ------------

def _replace_in_paragraph(paragraph, mapping):
    # Merge all runs, then rebuild – robust placeholder replacement
    if not paragraph.runs:
        paragraph.text = _replace_text(paragraph.text, mapping)
        return
    text = "".join(run.text for run in paragraph.runs)
    new_text = _replace_text(text, mapping)
    for _ in range(len(paragraph.runs) - 1):
        paragraph.runs[-1].clear()
        paragraph._p.remove(paragraph.runs[-1]._r)
    paragraph.runs[0].text = new_text

def _replace_text(text, mapping):
    if not text:
        return text
    for k, v in mapping.items():
        text = text.replace(k, v)
    return text

def fill_docx_template(template_path, mapping, out_path):
    doc = Document(template_path)
    # paragraphs
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_in_paragraph(p, mapping)
    doc.save(out_path)
    return out_path

# ------------ Quote builder (using your template) ------------

def build_quote_from_template(data: dict) -> str:
    # numbers / total formatting
    rate_text = str(data.get("rate", ""))
    qty_text = str(data.get("quantity", ""))
    total_text = str(data.get("total", ""))

    try:
        rate_num = float(str(rate_text).replace(",", "").replace("₹", ""))
        qty_num = float(str(qty_text).replace(",", ""))
        rate_text = f"₹{rate_num:,.2f}"
        total_text = f"₹{rate_num * qty_num:,.2f}"
    except Exception:
        pass

    mapping = {
        "{{Q_NO}}":          str(data.get("q_no", "")),
        "{{DATE}}":          str(data.get("date", datetime.utcnow().strftime("%Y-%m-%d"))),
        "{{COMPANY_NAME}}":  str(data.get("company_name", "")),
        "{{CUSTOMER_NAME}}": str(data.get("customer_name", "")),
        "{{PRODUCT}}":       str(data.get("product", "")),
        "{{QUANTITY}}":      str(qty_text),
        "{{UNITS}}":         str(data.get("units", "")),
        "{{RATE}}":          str(rate_text),
        "{{HSN}}":           str(data.get("hsn", "")),
        "{{TOTAL}}":         str(total_text),
    }

    template_path = os.path.join(os.path.dirname(__file__), "Template.docx")
    if not os.path.exists(template_path):
        raise RuntimeError("Template.docx not found in repo root")

    customer_safe = (data.get("customer_name") or "Customer").strip().replace(" ", "_")
    date_safe = (data.get("date") or datetime.utcnow().strftime("%Y-%m-%d"))
    out_path = f"/tmp/Quotation_{customer_safe}_{date_safe}.docx"

    return fill_docx_template(template_path, mapping, out_path)

# ------------ Gmail send ------------

def send_email_with_attachment(to_email: str, subject: str, body: str, attach_path: str):
    if not have_gmail_env():
        raise RuntimeError("Missing Gmail env vars (GMAIL_CLIENT_ID / SECRET / REFRESH_TOKEN / SENDER)")

    sender = os.environ["GMAIL_SENDER"]

    msg = MIMEMultipart()
    msg["To"] = to_email
    msg["From"] = sender
    msg["Subject"] = subject
    msg.attach(MIMEText(body, "plain"))

    with open(attach_path, "rb") as f:
        part = MIMEBase(
            "application",
            "vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        part.set_payload(f.read())
        encoders.encode_base64(part)
        part.add_header("Content-Disposition", f'attachment; filename="{os.path.basename(attach_path)}"')
    msg.attach(part)

    raw = base64.urlsafe_b64encode(msg.as_bytes()).decode("utf-8")
    service = gmail_service()
    try:
        sent = service.users().messages().send(userId="me", body={"raw": raw}).execute()
        return sent.get("id")
    except HttpError as e:
        raise RuntimeError(f"Gmail API error: {e}")

# ------------ Routes ------------

@app.get("/health")
def health():
    return jsonify({
        "service": "quotation-bot",
        "status": "ok",
        "gmail_ready": have_gmail_env(),
        "time": datetime.utcnow().isoformat() + "Z"
    })

@app.post("/send_quote")
def send_quote():
    """
    Body JSON example:
    {
      "q_no": "110",
      "date": "2025-11-06",
      "company_name": "NIVEE METAL PRODUCTS PVT LTD",
      "customer_name": "Rudra",
      "product": "5 inch SS 316L sheets",
      "quantity": "5",
      "rate": "25000",
      "units": "Pcs",
      "hsn": "7219",
      "email": "vip.vedant3@gmail.com"
    }
    """
    try:
        data = request.get_json(silent=True) or {}
        if not data.get("email"):
            return jsonify({"ok": False, "error": "Missing 'email'"}), 400

        # 1) Fill your Template.docx
        fp = build_quote_from_template(data)

        # 2) Email
        subject = f"Quotation {data.get('q_no','')} - {data.get('customer_name','')}".strip(" -")
        body = (
            f"Dear {data.get('customer_name','')},\n\n"
            f"Please find attached the quotation.\n\n"
            f"Regards,\n{data.get('company_name','')}"
        )
        msg_id = send_email_with_attachment(data["email"], subject or "Quotation", body, fp)

        return jsonify({"ok": True, "gmail_message_id": msg_id, "file": os.path.basename(fp)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=int(os.getenv("PORT", "8080")))
